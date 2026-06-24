from core.common.mongo.client import get_mongo_client
import pandas as pd
import subprocess
import os
import uuid
import re
import random
import boto3
import sys
from bson.objectid import ObjectId
from botocore.client import Config
import structlog
from pymongo import MongoClient
from datetime import datetime
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
PROJECT_ROOT = os.path.dirname(os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))))
sys.path.append(PROJECT_ROOT)
from constants import MinioConfig, MongoDBConfig, MigrateConfig, MongoDBCollectionConfig
from core.common.elastic import ElasticIndexer
from core.common.minio import MinIOClient
from logs.logger_conf import setup_logging

setup_logging()
logger = structlog.get_logger()

elastic_document = ElasticIndexer()
minio_client = MinIOClient()

s3 = boto3.client(
    "s3",
    endpoint_url=MinioConfig.ENDPOINT,
    aws_access_key_id=MinioConfig.ACCESS_KEY,
    aws_secret_access_key=MinioConfig.SECRET_KEY,
    config=Config(signature_version="s3v4"),
    region_name="us-east-1"
)

client = get_mongo_client()

db_raw = client[MigrateConfig.MIGRATE_RAW_DB]
signer_collection = db_raw['signer']
resource_collection = db_raw[MongoDBCollectionConfig.RAW_RESOURCE_COLLECTION_NAME]
document_segment_collection = db_raw[MongoDBCollectionConfig.RAW_DOCUMENTS_SEGMENTS_COLLECTION_NAME]

def format_date(date):
    """Convert date to string format YYYY-MM-DD HH:MM:SS or return empty string if None or invalid"""
    if pd.isna(date) or date is None:
        return ""
    try:
        # Nếu đã là datetime
        if isinstance(date, pd.Timestamp) or isinstance(date, datetime):
            return date.strftime("%Y-%m-%d %H:%M:%S")
        # Nếu là string -> parse về datetime
        parsed_date = pd.to_datetime(date, errors='coerce')
        if pd.isna(parsed_date):
            logger.warning(action="format_date", event="invalid_date_format", date_len=len(str(date)))
            return ""
        return parsed_date.strftime("%Y-%m-%d %H:%M:%S")
    except (ValueError, TypeError) as e:
        logger.error(action="format_date", event="format_date_failed", **{"error.code": "PARSE", "error.message": str(e)}, date_len=len(str(date)), exc_info=True)
        return ""



def gen_new_code():
    new_code = str(random.randint(100, 9999))
    while signer_collection.find_one({"code": new_code}):
        new_code = str(random.randint(100, 9999))
    return new_code

def gen_storage_code():
    # code = ulid.new().str
    # return "01" + code[2:]
    code =str(uuid.uuid4())
    return code


def get_or_create_signer_code(nguoi_ky, current_time):
    """Check if signers exist in signer_collection, create if not, and return list of codes"""
    if not nguoi_ky or pd.isna(nguoi_ky):
        return []
    
    # Split multiple signers by comma and clean names
    signers = [name.strip() for name in nguoi_ky.split(',') if name.strip()]
    signer_codes = []
    
    for signer_name in signers:
        # Normalize signer name for __text field (lowercase, no special chars)
        text_search = re.sub(r'\s+', '', signer_name.lower())
        
        # Check if signer exists
        signer = signer_collection.find_one({"name": signer_name})
        
        if signer:
            signer_codes.append(signer['code'])

        else:
            # Create new signer
            new_code = gen_new_code()
            signer_doc = {
                "_id": str(uuid.uuid4()),
                "code": new_code,
                "name": signer_name,
                "created_by": "",
                "created_date": current_time,
                "last_modified": current_time,
                "status": "active",
                "__text": f"{new_code} {text_search}"
            }
            
            try:
                signer_collection.insert_one(signer_doc)
                logger.info(action="get_or_create_signer_code", event="signer_added", signer_name=signer_name, code=new_code)
                signer_codes.append(new_code)
            except Exception as error:
                logger.error(action="get_or_create_signer_code", event="signer_addition_failed", **{"error.code": "DB", "error.message": str(error)}, signer_name=signer_name, exc_info=True)
                continue
    
    return signer_codes


def convert_ngay_hieu_luc(ngay_hieu_luc, ngay_ban_hanh):
    if ngay_hieu_luc == "Có hiệu lực từ ngày ban hành":
        return ngay_ban_hanh
    else:
        return ngay_hieu_luc


def list_objects():
    listDoc = []
    paginator = s3.get_paginator("list_objects_v2")

    for page in paginator.paginate(Bucket=MinioConfig.DEFAULT_BUCKET_NAME):
        for obj in page.get("Contents", []):
            # Lấy tên file và kích thước (bytes)
            file_info = {
                "name": obj["Key"],
                "size": obj["Size"]
            }
            listDoc.append(file_info)

    logger.info(action="list_objects", event="objects_listed", count=len(listDoc))
    return listDoc


def get_storage_code(doc_id, documents):
    match_doc = next((doc for doc in documents if str(doc_id) in doc.get("name", "")), None)
    if not match_doc:
        logger.error(action="get_storage_code", event="document_not_found", **{"error.code": "DB", "error.message": "Document not found"}, doc_id=doc_id)
        return ""

    doc_name = match_doc.get("name")
    doc_size = match_doc.get("size")
    existing_doc = resource_collection.find_one({"name": doc_name})
    if existing_doc:
        logger.info(action="get_storage_code", event="document_found", doc_id=doc_id, code=existing_doc['code'])
        return existing_doc['code']
    else:
        current_time = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        new_resource = {
            "_id": ObjectId(),
            "code": gen_storage_code(),  # Tạo code ngẫu nhiên
            "bucket": MinioConfig.DEFAULT_BUCKET_NAME,
            "created_by": "",
            "created_date": current_time,
            "last_modified": current_time,
            "last_modified_by": "",
            "name": doc_name,
            "path": doc_name,
            "size": doc_size,
            "status": "ACTIVE",
            "state": ""
        }
        try:
            resource_collection.insert_one(new_resource)
            logger.info(action="get_storage_code", event="resource_created", doc_name=doc_name, code=new_resource['code'])
            return new_resource['code']
        except Exception as e:
            logger.error(action="get_storage_code", event="resource_creation_failed", **{"error.code": "DB", "error.message": str(e)}, doc_name=doc_name, exc_info=True)
            return None

def convert_doc_to_docx(input_path, output_path=None):
    """
    Convert .doc -> .docx bằng LibreOffice
    input_path: đường dẫn file .doc
    output_path: nếu None thì lưu cùng thư mục với .doc
    """
    if not output_path:
        output_path = os.path.splitext(input_path)[0] + ".docx"

    try:
        subprocess.run([
            "soffice", "--headless", "--convert-to", "docx", "--outdir",
            os.path.dirname(output_path) or ".", input_path
        ], check=True)
        logger.info(action="convert_doc_to_docx", event="conversion_successful", output_path=output_path)
        return output_path
    except Exception as e:
        logger.error(action="convert_doc_to_docx", event="conversion_failed", **{"error.code": "EXT", "error.message": str(e)}, exc_info=True)
        return None

def get_doc_content(storage_code: str, object_name: str=None) -> str:
    temp_file_path = ""
    docx_path = ""
    
    try:
        # Tìm document với storage_code
        if object_name is None:
            document = resource_collection.find_one({"code": storage_code})
            if not document or "name" not in document:
                logger.warning(action="get_doc_content", event="storage_code_not_found", storage_code=storage_code)
                return ""
            name = document["name"]
            # object_name = f"{bucket}/uploads_record/{name}"
            object_name = f"{name}"
            logger.debug(action="get_doc_content", event="object_name_found", object_name=object_name)

        # Download file từ MinIO
        file_stream = minio_client.download_file(object_name=object_name)
        temp_file_path = f"temp_{name}"
        
        # Lưu file tạm thời
        with open(temp_file_path, "wb") as f:
            f.write(file_stream.read())
        logger.debug(action="get_doc_content", event="temp_file_saved", temp_file_path=temp_file_path)
        
        # Chuyển đổi .doc sang .docx
        docx_path = convert_doc_to_docx(temp_file_path)
        if not docx_path:
            logger.warning(action="get_doc_content", event="conversion_failed")
            os.remove(temp_file_path)  # Xóa file tạm thời nếu chuyển đổi thất bại
            return ""
        
        # Trích xuất nội dung từ .docx
        text_content = extract_docx_text(docx_path)
        if not text_content:
            logger.warning(action="get_doc_content", event="content_extraction_failed")
            os.remove(temp_file_path)  # Xóa file .doc
            os.remove(docx_path)      # Xóa file .docx
            return ""
        
        # Xóa các file tạm thời
        try:
            os.remove(temp_file_path)
            logger.debug(action="get_doc_content", event="temp_file_deleted", temp_file_path=temp_file_path)
            os.remove(docx_path)
            logger.debug(action="get_doc_content", event="docx_file_deleted", docx_path=docx_path)
        except Exception as e:
            logger.error(action="get_doc_content", event="temp_file_deletion_failed", **{"error.code": "IO", "error.message": str(e)}, exc_info=True)
        
        return text_content
    
    except Exception as e:
        logger.error(action="get_doc_content", event="content_extraction_process_failed", **{"error.code": "EXT", "error.message": str(e)}, exc_info=True)
        # Đảm bảo xóa file tạm thời nếu có lỗi
        if os.path.exists(temp_file_path):
            os.remove(temp_file_path)
            logger.debug(action="get_doc_content", event="temp_file_deleted_on_error", temp_file_path=temp_file_path)
        if docx_path and os.path.exists(docx_path):
            os.remove(docx_path)
            logger.debug(action="get_doc_content", event="docx_file_deleted_on_error", docx_path=docx_path)
        return ""

def extract_docx_text(docx_path):
    """
    Đọc nội dung file .docx và trả về text của các phần tử (paragraphs + tables),
    bỏ qua thành phần bị lỗi (vd: merge bảng gây lỗi grid_offset).
    """
    try:
        doc = Document(docx_path)
        text_elements = []

        for element in doc.element.body:
            try:
                # ---- Paragraph ----
                if element.tag == qn('w:p'):
                    para = next((p for p in doc.paragraphs if p._element == element), None)
                    if para and para.text.strip():
                        text_elements.append(para.text.strip())

                # ---- Table ----
                elif element.tag == qn('w:tbl'):
                    table = next((t for t in doc.tables if t._tbl == element), None)
                    if not table:
                        continue

                    for r_idx, row in enumerate(table.rows):
                        try:
                            for c_idx, cell in enumerate(row.cells):
                                if cell.text.strip():
                                    text_elements.append(cell.text.strip())
                        except Exception as e:
                            logger.error(action="extract_docx_text", event="cell_processing_failed", **{"error.code": "PARSE", "error.message": str(e)}, row_idx=r_idx, exc_info=True)
                            continue

            except Exception as e:
                logger.error(action="extract_docx_text", event="element_processing_failed", **{"error.code": "PARSE", "error.message": str(e)}, exc_info=True)
                continue

        return "\n".join(text_elements)

    except Exception as e:
        logger.error(action="extract_docx_text", event="docx_extraction_failed", **{"error.code": "IO", "error.message": str(e)}, exc_info=True)
        return ""


def get_decree_status(tinh_trang):
    if "Còn hiệu lực đến" in tinh_trang:
        date_expired = tinh_trang.split(":")[1].strip()
        decree_status = tinh_trang.split("đến")[0].strip()
        return decree_status, date_expired
    elif "Đã có hiệu lực" in tinh_trang:
        return "Còn hiệu lực", ""
    else:
        return tinh_trang.strip(), "" 

if __name__ == "__main__":
    # connect_postgres()
    # list_objects()
    storage_code = "74441224-237e-48c9-a7e7-0b678cef47ce"
    logger.info(action="main", event="content_length_logged", content_len=len(get_doc_content(storage_code)))