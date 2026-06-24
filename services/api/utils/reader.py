import structlog
from docx import Document
from uuid import uuid4
from pathlib import Path
import sys
import os
import tempfile
import subprocess
import io
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import html
PROJECT_ROOT = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
sys.path.append(PROJECT_ROOT)
from logs.logger_conf import setup_logging

setup_logging()
logger = structlog.get_logger()


def validate_upload_file(file) -> str:    
    # Check file format
    if not (file.filename.endswith('.doc') or file.filename.endswith('.docx')):
        logger.warning("validate_upload_file_invalid_format", action="validate_upload_file", filename=file.filename)
        raise ValueError("Invalid file format. Only .doc and .docx are allowed")

    # Check file size (e.g., max 20MB)
    file.seek(0, os.SEEK_END)
    file_size = file.tell()
    if file_size > 20 * 1024 * 1024:
        logger.warning("validate_upload_file_too_large", action="validate_upload_file", filename=file.filename, size=file_size)
        raise ValueError("File size exceeds 20MB")
    file.seek(0)    


def read_file_docx(path_file: str) -> str:
    """Read a docx file and return a document"""
    try:
        if isinstance(path_file, str):
            if not Path(path_file).exists():
                raise FileNotFoundError(f"Could not find file: {path_file}")
            logger.debug("read_file_docx_started", action="read_file_docx", path=path_file)
            docx_document = Document(path_file)
            doc_content = "\n\n".join([para.text for para in docx_document.paragraphs])            
        else:
            raise ValueError("path_file must be a string")
        return doc_content    
    except Exception as e:
        logger.error("read_file_docx_failed", action="read_file_docx", **{"error.code": "IO", "error.message": str(e)}, exc_info=True)
        return None


def convert_doc_to_docx(file_path: str) -> str:
    """
    Convert a .doc file to .docx if necessary.

    Args:
        file_path (str): Path to the input file.

    Returns:
        str: Path to the converted .docx file if conversion succeeds, None otherwise.
    """
    try:
        file_name = os.path.basename(file_path)
        file_ext = os.path.splitext(file_name)[1].lower()

        if file_ext == '.doc':
            # Create a temporary directory for conversion
            path_folder = os.path.dirname(file_path)
            # Convert .doc to .docx using LibreOffice
            output_file_name = os.path.splitext(file_name)[0] + '.docx'
            output_file_path = os.path.join(path_folder, output_file_name)

            cmd = [
                'soffice',
                '--headless',
                '--convert-to', 'docx',
                '--outdir', path_folder,
                file_path
            ]
            result = subprocess.run(cmd, capture_output=True, text=True)
            if result.returncode != 0:
                logger.error("convert_doc_to_docx_conversion_failed", action="convert_doc_to_docx", **{"error.code": "IO", "error.message": "LibreOffice conversion failed"}, stderr_len=len(result.stderr or ""))
                return None

            if not os.path.exists(output_file_path):
                logger.error("convert_doc_to_docx_output_not_found", action="convert_doc_to_docx", **{"error.code": "IO", "error.message": "Converted file not found"}, path=output_file_path)
                return None

            return output_file_path
        else:
            return file_path
    except Exception as e:
        logger.error("convert_doc_to_docx_failed", action="convert_doc_to_docx", **{"error.code": "IO", "error.message": str(e)}, exc_info=True)
        return None

def docx_to_html(docx_input):
    """
    Convert a .docx file to HTML while preserving formatting.

    Args:
        docx_input: Path to .docx file (str) or file-like object (e.g., BytesIO).

    Returns:
        str: HTML string representing the .docx content with formatting preserved.

    Raises:
        ValueError: If docx_input is neither a string nor BytesIO.
        Exception: For errors reading or processing the .docx file.
    """
    # Load the .docx file
    try:
        if isinstance(docx_input, str):
            doc = Document(docx_input)
        elif isinstance(docx_input, io.BytesIO):
            doc = Document(docx_input)
        else:
            raise ValueError("docx_input must be a file path (str) or BytesIO object")
    except Exception as e:
        raise Exception(f"Failed to load .docx file: {str(e)}")

    # Initialize HTML content
    html_parts = []
    current_list = None  # Track open list type and level (e.g., {"tag": "ul", "level": 0})

    # Process each paragraph
    for para in doc.paragraphs:
        if not para.text.strip() and not para.runs:
            continue  # Skip empty paragraphs

        # Determine paragraph style and alignment
        style_name = para.style.name.lower()
        alignment = para.alignment
        align_style = ""
        if alignment == WD_ALIGN_PARAGRAPH.CENTER:
            align_style = "text-align: center;"
        elif alignment == WD_ALIGN_PARAGRAPH.RIGHT:
            align_style = "text-align: right;"
        elif alignment == WD_ALIGN_PARAGRAPH.JUSTIFY:
            align_style = "text-align: justify;"
        else:
            align_style = "text-align: left;"

        # Initialize paragraph content
        para_content = []

        # Process each run in the paragraph
        for run in para.runs:
            text = html.escape(run.text)
            if not text:
                continue

            # Build run styles
            run_styles = []
            if run.bold:
                text = f"<strong>{text}</strong>"
            if run.italic:
                text = f"<em>{text}</em>"
            if run.underline:
                text = f"<u>{text}</u>"

            # Font properties
            font = run.font
            if font.name:
                run_styles.append(f"font-family: '{font.name}';")
            if font.size:
                font_size_pt = font.size.pt
                run_styles.append(f"font-size: {font_size_pt}pt;")
            if font.color.rgb:
                color = f"#{font.color.rgb:06x}"
                run_styles.append(f"color: {color};")

            # Combine run styles
            style_attr = f' style="{"".join(run_styles)}"' if run_styles else ""
            para_content.append(f"<span{style_attr}>{text}</span>")

        # Combine paragraph content
        para_html = ''.join(para_content)

        # Handle lists
        if para._element.xpath(".//w:numPr"):
            num_id = para._element.find(".//w:numId", namespaces={"w": qn("w:numId")})
            ilvl = para._element.find(".//w:ilvl", namespaces={"w": qn("w:ilvl")})
            level = int(ilvl.get(qn("w:val"))) if ilvl is not None else 0
            is_bullet = style_name.startswith("list bullet")
            list_tag = "ul" if is_bullet else "ol"

            # Close lists if level decreases or type changes
            if current_list:
                current_tag, current_level = current_list["tag"], current_list["level"]
                if level < current_level or (list_tag != current_tag and level <= current_level):
                    for _ in range(current_level, level - 1, -1):
                        html_parts.append(f"</{current_tag}>")
                    current_list = None

            # Open new list if needed
            if not current_list or current_list["tag"] != list_tag or current_list["level"] != level:
                if current_list:
                    html_parts.append(f"</{current_list['tag']}>")
                html_parts.append(f'<{list_tag} style="margin-left: {20 * (level + 1)}px;">')
                current_list = {"tag": list_tag, "level": level}

            html_parts.append(f'<li style="{align_style}">{para_html}</li>')
        else:
            # Close any open list
            if current_list:
                html_parts.append(f"</{current_list['tag']}>")
                current_list = None

            # Map style to HTML tag
            if "heading" in style_name:
                level = min(int(style_name.replace("heading", "").strip() or 1), 6)
                html_parts.append(f'<h{level} style="{align_style}">{para_html}</h{level}>')
            else:
                html_parts.append(f'<p style="{align_style}">{para_html}</p>\n')

    # Process tables
    for table in doc.tables:
        # Close any open list
        if current_list:
            html_parts.append(f"</{current_list['tag']}>")
            current_list = None

        html_parts.append('<table border="1" style="border-collapse: collapse; width: 100%;">\n')
        for row in table.rows:
            html_parts.append("<tr>")
            for cell in row.cells:
                cell_content = []
                for para in cell.paragraphs:
                    if not para.text.strip():
                        continue
                    # Build run content for each paragraph in the cell
                    para_content = []
                    for run in para.runs:
                        if not run.text:
                            continue
                        text = html.escape(run.text)
                        # Apply inline formatting
                        if run.bold:
                            text = f"<strong>{text}</strong>"
                        if run.italic:
                            text = f"<em>{text}</em>"
                        if run.underline:
                            text = f"<u>{text}</u>"

                        # Build style attributes
                        run_styles = []
                        if run.font.name:
                            run_styles.append(f"font-family: '{run.font.name}';")
                        if run.font.size:
                            run_styles.append(f"font-size: {run.font.size.pt}pt;")
                        if run.font.color.rgb:
                            run_styles.append(f"color: #{run.font.color.rgb:06x};")

                        style_attr = f' style="{"".join(run_styles)}"' if run_styles else ""
                        para_content.append(f"<span{style_attr}>{text}</span>")

                    cell_content.append(f"<p>{''.join(para_content)}</p>")
                html_parts.append(f'<td style="padding: 5px;">{"".join(cell_content)}</td>')
            html_parts.append("</tr>\n")
        html_parts.append("</table>\n")

    # Close any remaining open list
    if current_list:
        html_parts.append(f"</{current_list['tag']}>")

    # Build final HTML
    html_output = f"""
    <!DOCTYPE html>
    <html lang="en">
    <head>
        <meta charset="UTF-8">
        <meta name="viewport" content="width=device-width, initial-scale=1.0">
        <title>Document</title>
        <style>
            body {{
                font-family: Arial, sans-serif;
                max-width: 800px;
                margin: 20px auto;
                line-height: 1.6;
                color: #333;
            }}
            h1, h2, h3, h4, h5, h6 {{
                color: #2c3e50;
                margin: 10px 0;
            }}
            p {{
                margin: 8px 0;
            }}
            ul, ol {{
                margin: 10px 0;
                padding-left: 20px;
            }}
            li {{
                margin-bottom: 5px;
            }}
            table {{
                border-collapse: collapse;
                width: 100%;
                margin: 10px 0;
            }}
            td, th {{
                border: 1px solid #ddd;
                padding: 8px;
            }}
        </style>
    </head>
    <body>
    {"".join(html_parts)}
    </body>
    </html>
    """

    return html_output

import mammoth
def convert_docx_to_html(docx_file_path):
    result = mammoth.convert_to_html(docx_file_path)    
    return result.value     

HTML_TEMPLATE = """<!DOCTYPE html PUBLIC \"-//W3C//DTD XHTML 1.0 Transitional//EN\"
    \"http://www.w3.org/TR/xhtml1/DTD/xhtml1-transitional.dtd\">
<html xmlns=\"http://www.w3.org/1999/xhtml\">

<head>
    <meta http-equiv=\"Content-Type\" content=\"text/html; charset=utf-8\" />
    <title></title>
    <style type=\"text/css\">
        .csBEC7F46E {}

        .cs15A7E4B {
            width: 30%;
            padding: 2.85pt 5.4pt 2.85pt 5.4pt;
            border-top: none;
            border-right: none;
            border-bottom: none;
            border-left: none
        }

        .cs947971CC {
            text-align: center;
            text-indent: 0pt;
            margin: 0pt 0pt 0pt 0pt
        }

        .csE58370F {
            color: #000000;
            background-color: transparent;
            font-family: 'DejaVu Serif';
            font-size: 12pt;
            font-weight: bold;
            font-style: normal;
        }

        .cs636E4293 {
            width: 66%;
            padding: 2.85pt 5.4pt 2.85pt 5.4pt;
            border-top: none;
            border-right: none;
            border-bottom: none;
            border-left: none
        }

        .csF6C3AA98 {
            color: #000000;
            background-color: transparent;
            font-family: 'DejaVu Serif';
            font-size: 12pt;
            font-weight: normal;
            font-style: normal;
        }

        .csB8450DD5 {
            text-align: right;
            text-indent: 0pt;
            margin: 0pt 0pt 0pt 0pt
        }

        .cs3805C34 {
            color: #000000;
            background-color: transparent;
            font-family: 'DejaVu Serif';
            font-size: 12pt;
            font-weight: normal;
            font-style: italic;
        }

        .csA81C8FD {
            text-align: center;
            text-indent: 0pt;
            margin: 0pt 0pt 6pt 0pt
        }

        .csE6DAD51 {
            text-align: left;
            text-indent: 0pt;
            margin: 0pt 0pt 6pt 0pt
        }

        .cs2C4BD76E {
            color: #000000;
            background-color: transparent;
            font-family: 'DejaVu Serif';
            font-size: 12pt;
            font-weight: normal;
            font-style: normal;
            text-decoration: line-through;
        }

        .csF426EBD5 {
            width: 211.95pt;
            padding: 0pt 5.4pt 0pt 5.4pt;
            border-top: none;
            border-right: none;
            border-bottom: none;
            border-left: none
        }

        .cs338ED9EC {
            text-align: left;
            text-indent: 0pt;
            margin: 0pt 0pt 0pt 0pt
        }

        .cs47C5D3E0 {
            width: 214.15pt;
            padding: 0pt 5.4pt 0pt 5.4pt;
            border-top: none;
            border-right: none;
            border-bottom: none;
            border-left: none
        }

        .cs78F37DAA {
            width: 30.2pt;
            padding: 1.4pt 5.4pt 1.4pt 5.4pt;
            border-top: 1pt windowtext solid;
            border-right: 1pt windowtext solid;
            border-bottom: 1pt windowtext solid;
            border-left: 1pt windowtext solid
        }

        .csFFEBCAEA {
            width: 92.25pt;
            padding: 1.4pt 5.4pt 1.4pt 5.4pt;
            border-top: 1pt windowtext solid;
            border-right: 1pt windowtext solid;
            border-bottom: 1pt windowtext solid;
            border-left: none
        }

        .cs88A329F1 {
            width: 157.5pt;
            padding: 1.4pt 5.4pt 1.4pt 5.4pt;
            border-top: 1pt windowtext solid;
            border-right: 1pt windowtext solid;
            border-bottom: 1pt windowtext solid;
            border-left: none
        }

        .cs338AE7BD {
            width: 140.25pt;
            padding: 1.4pt 5.4pt 1.4pt 5.4pt;
            border-top: 1pt windowtext solid;
            border-right: 1pt windowtext solid;
            border-bottom: 1pt windowtext solid;
            border-left: none
        }

        .csE6811EE2 {
            width: 30.2pt;
            padding: 1.4pt 5.4pt 1.4pt 5.4pt;
            border-top: none;
            border-right: 1pt windowtext solid;
            border-bottom: 1pt windowtext solid;
            border-left: 1pt windowtext solid
        }

        .cs3D548C61 {
            width: 92.25pt;
            padding: 1.4pt 5.4pt 1.4pt 5.4pt;
            border-top: none;
            border-right: 1pt windowtext solid;
            border-bottom: 1pt windowtext solid;
            border-left: none
        }

        .cs521B3205 {
            width: 157.5pt;
            padding: 1.4pt 5.4pt 1.4pt 5.4pt;
            border-top: none;
            border-right: 1pt windowtext solid;
            border-bottom: 1pt windowtext solid;
            border-left: none
        }

        .cs41A9B2C5 {
            width: 140.25pt;
            padding: 1.4pt 5.4pt 1.4pt 5.4pt;
            border-top: none;
            border-right: 1pt windowtext solid;
            border-bottom: 1pt windowtext solid;
            border-left: none
        }
    </style>
</head>

<body>
    {html_content}
</body>
</html>"""

# Example usage
if __name__ == "__main__":
    # Example file to upload
    input_file = "/home/ubuntu/projects/AI/git/users/giangnv/law-document-sync-core-service/services/uploads/DOWNLOAD_FILE.doc"
    new_file = convert_doc_to_docx(input_file)
    if new_file:
        logger.info("process_and_upload_success", action="main", file=new_file)
        convert_docx_to_html(new_file)
    else:
        logger.error("process_and_upload_failed", action="main", **{"error.code": "IO", "error.message": "Processing failed"})