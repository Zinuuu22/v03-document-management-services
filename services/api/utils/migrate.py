from core.common.mongo.client import get_mongo_client
import uuid
import os
import sys
from io import BytesIO
from docx import Document
import requests
from pymongo import MongoClient
import structlog

PROJECT_ROOT = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
sys.path.append(PROJECT_ROOT)
from constants import MongoDBConfig, APIEndpoints, MongoDBCollectionConfig, MigrateConfig
from logs.logger_conf import setup_logging

setup_logging()
logger = structlog.get_logger()


client = get_mongo_client()
db = client[MigrateConfig.MIGRATE_CORE_DB]
resource_collection = db[MongoDBCollectionConfig.LAW_DOCUMENT_STORAGE_COLLECTION_NAME]

IMAGE_FOLDER = os.path.join(PROJECT_ROOT, "./images")
os.makedirs(IMAGE_FOLDER, exist_ok=True)


def download_docx(code, name, max_retries=2):        
    url = f"{APIEndpoints.DOWNLOAD_DOCX}/{code}"
    headers = {"accept": "*/*"}
    retries = 0

    logger.debug("download_docx_started", action="download_docx", name=name, code=code)
    while retries <= max_retries:
        try:
            headers = {"accept": "*/*"}
            response = requests.post(url, headers=headers, data="")
            response.raise_for_status()
            source_stream = BytesIO(response.content)
            document = Document(source_stream)
            logger.debug("download_docx_success", action="download_docx", name=name)
            return document
        except Exception as e:
            retries += 1
            logger.error("download_docx_failed", action="download_docx", **{"error.code": "EXT", "error.message": str(e)}, name=name, retry=retries, exc_info=True)
            if retries > max_retries:
                logger.error("download_docx_max_retries_exceeded", action="download_docx", **{"error.code": "EXT", "error.message": "Max retries reached"}, name=name, max_retries=max_retries)
                return None


def clean_filename(title):
    """Clean filename to remove invalid characters."""
    return "".join(c for c in title if c.isalnum() or c in (" ", "_")).strip()

def extract_content_from_docx(document, document_id=None, title='No Name', code=None):
    logger.debug("extract_content_from_docx_started", action="extract_content_from_docx", title=title, document_id=document_id)
    try:
        elements = []
        position = 1
        temp_paragraph = []
        image_index = 0
        cleaned_title = clean_filename(title)  # Compute once

        # Ensure image folder exists
        os.makedirs(IMAGE_FOLDER, exist_ok=True)

        def save_paragraph():
            """Combine and save consecutive paragraphs."""
            nonlocal position
            if temp_paragraph:
                combined_text = "\n".join(temp_paragraph).strip()
                if combined_text:  # Only save non-empty paragraphs
                    elements.append({
                        "document_id": document_id,
                        "code": code,
                        "title": cleaned_title,
                        "type": "paragraph",
                        "content": combined_text,
                        "position": position,
                    })
                    position += 1
                temp_paragraph.clear()  # Clear immediately to free memory

        def save_image(paragraph, img_index):
            """Extract and save images from a paragraph."""
            nonlocal position
            for run in paragraph.runs:
                for elem in run._element.iter():
                    if elem.tag.endswith("blip"):
                        rel_id = elem.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
                        if rel_id in document.part.related_parts:
                            image_part = document.part.related_parts[rel_id]
                            image_path = os.path.join(IMAGE_FOLDER, f"{cleaned_title}_{img_index + 1}.png")
                            with open(image_path, "wb") as img_file:
                                img_file.write(image_part.blob)
                            image_id = str(uuid.uuid4())
                            elements.append({
                                "document_id": document_id,
                                "code": code,
                                "title": cleaned_title,
                                "type": "image",
                                "image_id": image_id,
                                "image_path": image_path,
                                "description": f"{cleaned_title}_{img_index + 1}",
                                "position": position,
                            })
                            position += 1
                            return True
            return False

        def save_table(table, tbl_index):
            """Extract and save data from a table."""
            nonlocal position
            table_id = str(uuid.uuid4())
            table_data = {
                "document_id": document_id,
                "code": code,
                "table_id": table_id,
                "title": cleaned_title,
                "type": "table",
                "position": position,
                "rows": [],
            }
            # Optional: Extract table content if needed
            # for row in table.rows:
            #     row_data = [cell.text.strip() for cell in row.cells]
            #     table_data["rows"].append(row_data)
            elements.append(table_data)
            position += 1

        # Process paragraphs and images
        for paragraph in document.paragraphs:
            paragraph_text = paragraph.text.strip()
            if paragraph_text:
                temp_paragraph.append(paragraph_text)

            # Check for images in the paragraph
            if save_image(paragraph, image_index):
                save_paragraph()  # Save any accumulated text before image
                image_index += 1

        # Save any remaining paragraphs
        save_paragraph()

        # Process tables
        for table in document.tables:
            save_paragraph()  # Save any accumulated text before table
            save_table(table, image_index)

            logger.debug("extract_content_from_docx_success", action="extract_content_from_docx", title=title)
        return elements

    except Exception as e:
        logger.error("extract_content_from_docx_failed", action="extract_content_from_docx", **{"error.code": "IO", "error.message": str(e)}, title=title, exc_info=True)
        return []


def post_process(element):
    '''
        format elements and convert noise elements to masks
    '''
    _elements_text = []
    for _element in element:
        _element_type = _element['type']
        if _element_type == 'table':
            table_id = _element['table_id']    
            _elements_text.append(f'REFERENCE_TABLE_{table_id}')
        elif _element_type == 'paragraph':
            paragraph_content = _element['content']
            _elements_text.append(paragraph_content)
        elif _element_type == 'image':
            image_id = _element['image_id']    
            _elements_text.append(f'REFERENCE_IMAGE_{image_id}')
    return '\n'.join(_elements_text)


def preprocess_document_from_storage_code(storage_code):
    '''
        read file from 'storage_code' and preprocess file content
    '''

    res_document = resource_collection.find_one({'storage_id': storage_code})
    if res_document is None:
        logger.warning("preprocess_document_storage_code_not_found", action="preprocess_document_from_storage_code", storage_code=storage_code)
        raise Exception(f"ERROR: NOT FOUND {storage_code}")
        
    try:
        document_id, name, path, code = res_document["_id"], res_document["name"], res_document.get("path", ""), res_document["storage_id"]
    except Exception as e:
        logger.error("preprocess_document_extract_fields_failed", action="preprocess_document_from_storage_code", **{"error.code": "DB", "error.message": str(e)}, storage_code=storage_code, exc_info=True)
        raise Exception(f"EXTRACT PATH, NAME, CODE ERROR: {e}")
        
    if str(name).lower().find('pdf') != -1:
        raise Exception("Error: Can not process pdf file")

    docx_document = download_docx(code, name)
    if not docx_document:
        logger.warning("preprocess_document_download_failed", action="preprocess_document_from_storage_code", code=code)
        raise Exception(f"Failed Download File with {code}")
    else:
        logger.debug("preprocess_document_download_success", action="preprocess_document_from_storage_code", code=code)
        try:
            elements = extract_content_from_docx(docx_document, document_id, name, code)
        except Exception as e:
            logger.error("preprocess_document_extract_content_failed", action="preprocess_document_from_storage_code", **{"error.code": "IO", "error.message": str(e)}, exc_info=True)
            raise Exception(f"extract_data_from_docx(): {e}")
        
        logger.debug("preprocess_document_extract_content_success", action="preprocess_document_from_storage_code", element_count=len(elements))
        try:
            text = post_process(elements)
        except Exception as e:
            logger.error("preprocess_document_postprocess_failed", action="preprocess_document_from_storage_code", **{"error.code": "PARSE", "error.message": str(e)}, exc_info=True)
            raise Exception(f"Error in post_process(): {e}")
        return text
    

def preprocess_document_from_stream(source_stream):            
    '''
        read file from 'source_stream' and preprocess file content
    '''
    docx_document = Document(source_stream)
            
    if not docx_document:
        logger.warning("preprocess_document_stream_load_failed", action="preprocess_document_from_stream")
        raise Exception("Failed Load File from Stream")
    else:
        logger.debug("preprocess_document_stream_load_success", action="preprocess_document_from_stream")
        try:
            elements = extract_content_from_docx(docx_document)
        except Exception as e:
            logger.error("preprocess_document_extract_content_failed", action="preprocess_document_from_stream", **{"error.code": "IO", "error.message": str(e)}, exc_info=True)
            raise Exception(f"extract_data_from_docx(): {e}")
        
        logger.debug("preprocess_document_extract_content_success", action="preprocess_document_from_stream", element_count=len(elements))
        try:
            text = post_process(elements)
        except Exception as e:
            logger.error("preprocess_document_postprocess_failed", action="preprocess_document_from_stream", **{"error.code": "PARSE", "error.message": str(e)}, exc_info=True)
            raise Exception(f"Error in post_process(): {e}")
        return text