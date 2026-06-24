import html
import os
import subprocess
import sys
import tempfile
import time
from io import BufferedReader, BytesIO
from pathlib import Path
from typing import Optional, Union

import mammoth
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from werkzeug.datastructures import FileStorage

PROJECT_ROOT = os.path.dirname(os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))))
sys.path.append(PROJECT_ROOT)

import structlog
from logs.logger_conf import setup_logging

setup_logging()
logger = structlog.get_logger()

# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------

MAX_FILE_SIZE_BYTES = 20 * 1024 * 1024  # 20 MB
ALLOWED_EXTENSIONS  = {".doc", ".docx"}

_MAMMOTH_HTML_TEMPLATE = """<!DOCTYPE html PUBLIC "-//W3C//DTD XHTML 1.0 Transitional//EN"
    "http://www.w3.org/TR/xhtml1/DTD/xhtml1-transitional.dtd">
<html xmlns="http://www.w3.org/1999/xhtml">
<head>
    <meta http-equiv="Content-Type" content="text/html; charset=utf-8" />
    <title>Document</title>
    <style type="text/css">
        body {{ font-family: 'DejaVu Serif', serif; font-size: 12pt; color: #000; }}
        table {{ border-collapse: collapse; width: 100%; margin: 10px 0; }}
        td, th {{ border: 1px solid #000; padding: 4px 8px; }}
    </style>
</head>
<body>
{html_content}
</body>
</html>"""


# ---------------------------------------------------------------------------
# DocumentProcessor
# ---------------------------------------------------------------------------

class DocumentProcessor:
    """Xử lý file tài liệu (.doc, .docx): validate, đọc nội dung, chuyển đổi HTML."""

    # ------------------------------------------------------------------
    # Validate
    # ------------------------------------------------------------------

    def validate_file(self, file: Union[FileStorage, str]) -> None:
        """
        Kiểm tra định dạng và kích thước file.

        Args:
            file: FileStorage từ Flask request hoặc đường dẫn file (str).

        Raises:
            ValueError: Nếu định dạng không hợp lệ hoặc kích thước vượt quá 20MB.
            FileNotFoundError: Nếu đường dẫn file không tồn tại.
        """
        if isinstance(file, str):
            path = Path(file)
            if not path.exists():
                raise FileNotFoundError(f"File not found: {file}")
            file_name = path.name
            file_size = path.stat().st_size
            _check_extension(file_name)
            _check_size(file_name, file_size)

        elif isinstance(file, FileStorage):
            file_name = file.filename or ""
            _check_extension(file_name)
            file.stream.seek(0, os.SEEK_END)
            file_size = file.stream.tell()
            file.stream.seek(0)
            _check_size(file_name, file_size)

        else:
            raise ValueError("file must be FileStorage or str")

        logger.info("validate_file_success", action="validate_file", filename=file_name)

    # ------------------------------------------------------------------
    # Read .docx
    # ------------------------------------------------------------------

    def read_docx(self, docx_input: Union[str, BufferedReader, FileStorage]) -> Optional[str]:
        start_t = time.time()
        try:
            doc = _load_docx(docx_input)
            content = []

            body = doc.element.body

            for child in body.iterchildren():
                tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag

                if tag == 'p':
                    from docx.text.paragraph import Paragraph
                    para = Paragraph(child, doc)
                    if para.text.strip():
                        content.append(para.text)

                elif tag == 'tbl':
                    from docx.table import Table
                    table = Table(child, doc)
                    for row in table.rows:
                        row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                        if row_text:
                            content.append(" | ".join(row_text))

            result = "\n\n".join(content)
            logger.info("read_docx_success", action="read_docx", **{"event.duration": time.time()-start_t, "event.status": "success"}, chars=len(result))
            return result

        except Exception as e:
            logger.error("read_docx_failed", action="read_docx", **{"error.code": "IO", "error.message": str(e), "event.duration": time.time()-start_t, "event.status": "failure"}, exc_info=True)
            return None
    # ------------------------------------------------------------------
    # Convert .doc → .docx
    # ------------------------------------------------------------------

    def convert_doc_to_docx(self, file_input: Union[str, FileStorage, BytesIO]) -> Optional[str]:
        """
        Chuyển đổi file .doc sang .docx bằng LibreOffice.
        Nếu file đã là .docx thì trả về đường dẫn gốc.

        Args:
            file_input: Đường dẫn file (str), FileStorage, hoặc BytesIO.

        Returns:
            Đường dẫn tới file .docx, hoặc None nếu thất bại.
        """
        temp_input_path  = None
        start_t = time.time()
        try:
            # Ghi vào temp file nếu cần
            if isinstance(file_input, FileStorage):
                suffix = Path(file_input.filename or "file.doc").suffix.lower()
                temp_input_path = _write_temp_file(file_input.stream.read(), suffix)
                file_path = temp_input_path

            elif isinstance(file_input, BytesIO):
                temp_input_path = _write_temp_file(file_input.read(), ".doc")
                file_path = temp_input_path

            elif isinstance(file_input, str):
                if not Path(file_input).exists():
                    raise FileNotFoundError(f"File not found: {file_input}")
                file_path = file_input

            else:
                raise ValueError("file_input must be str, FileStorage, or BytesIO")

            if Path(file_path).suffix.lower() != ".doc":
                logger.info("convert_doc_to_docx_success", action="convert_doc_to_docx", **{"event.duration": time.time()-start_t, "event.status": "success"}, path=file_path, note="no_conversion_needed")
                return file_path

            # Chạy LibreOffice để convert
            out_dir = os.path.dirname(file_path) or tempfile.gettempdir()
            cmd = ["soffice", "--headless", "--convert-to", "docx", "--outdir", out_dir, file_path]
            result = subprocess.run(cmd, capture_output=True, text=True)

            if result.returncode != 0:
                logger.error("convert_doc_to_docx_failed", action="convert_doc_to_docx", **{"error.code": "IO", "error.message": result.stderr, "event.duration": time.time()-start_t, "event.status": "failure"})
                return None

            output_path = os.path.join(out_dir, Path(file_path).stem + ".docx")
            if not os.path.exists(output_path):
                logger.error("convert_doc_to_docx_failed", action="convert_doc_to_docx", **{"error.code": "IO", "error.message": "converted_file_missing", "event.duration": time.time()-start_t, "event.status": "failure"}, path=output_path)
                return None

            logger.info("convert_doc_to_docx_success", action="convert_doc_to_docx", **{"event.duration": time.time()-start_t, "event.status": "success"}, source=file_path, target=output_path)
            return output_path

        except Exception as e:
            logger.error("convert_doc_to_docx_failed", action="convert_doc_to_docx", **{"error.code": "IO", "error.message": str(e), "event.duration": time.time()-start_t, "event.status": "failure"}, exc_info=True)
            return None

        finally:
            # Chỉ xoá temp input, giữ lại output để caller dùng
            if temp_input_path and os.path.exists(temp_input_path):
                os.remove(temp_input_path)

    # ------------------------------------------------------------------
    # Convert .docx → HTML (manual, preserve formatting)
    # ------------------------------------------------------------------

    def docx_to_html(self, docx_input: Union[str, BytesIO]) -> str:
        """
        Chuyển đổi .docx sang HTML, giữ nguyên định dạng (bold, italic, heading, table, list).

        Args:
            docx_input: Đường dẫn file (str) hoặc BytesIO.

        Returns:
            Chuỗi HTML đầy đủ.

        Raises:
            FileNotFoundError: Nếu file không tồn tại.
            ValueError: Nếu input không hợp lệ.
        """
        start_t = time.time()
        try:
            doc = _load_docx(docx_input)
            html_parts   = []
            current_list = None  # {"tag": "ul"|"ol", "level": int}

            for para in doc.paragraphs:
                if not para.text.strip() and not para.runs:
                    continue

                align_style = _get_alignment_style(para)
                para_html   = _render_runs(para.runs)

                # --- List handling ---
                num_pr = para._element.xpath(".//w:numPr", namespaces={"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"})
                if num_pr:
                    ilvl_el = num_pr[0].find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}ilvl")
                    level   = int(ilvl_el.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val", "0")) if ilvl_el is not None else 0
                    list_tag = "ul" if para.style.name.lower().startswith("list bullet") else "ol"

                    if current_list and (current_list["tag"] != list_tag or current_list["level"] != level):
                        html_parts.append(f"</{current_list['tag']}>")
                        current_list = None

                    if not current_list:
                        html_parts.append(f'<{list_tag} style="margin-left: {20 * (level + 1)}px;">')
                        current_list = {"tag": list_tag, "level": level}

                    html_parts.append(f'<li style="{align_style}">{para_html}</li>')

                else:
                    # Đóng list nếu đang mở
                    if current_list:
                        html_parts.append(f"</{current_list['tag']}>")
                        current_list = None

                    style_name = para.style.name.lower()
                    if "heading" in style_name:
                        level = min(int("".join(filter(str.isdigit, style_name)) or "1"), 6)
                        html_parts.append(f'<h{level} style="{align_style}">{para_html}</h{level}>')
                    else:
                        html_parts.append(f'<p style="{align_style}">{para_html}</p>')

            # --- Tables ---
            for table in doc.tables:
                if current_list:
                    html_parts.append(f"</{current_list['tag']}>")
                    current_list = None
                html_parts.append(_render_table(table))

            if current_list:
                html_parts.append(f"</{current_list['tag']}>")

            body = "\n".join(html_parts)
            output = _wrap_html(body)
            logger.info("docx_to_html_success", action="docx_to_html", **{"event.duration": time.time()-start_t, "event.status": "success"}, chars=len(output))
            return output

        except Exception as e:
            logger.error("docx_to_html_failed", action="docx_to_html", **{"error.code": "IO", "error.message": str(e), "event.duration": time.time()-start_t, "event.status": "failure"}, exc_info=True)
            raise

    # ------------------------------------------------------------------
    # Convert .docx → HTML (mammoth, simpler)
    # ------------------------------------------------------------------

    def docx_to_html_mammoth(self, docx_input: Union[str, BytesIO]) -> str:
        """
        Chuyển đổi .docx sang HTML bằng mammoth (đơn giản hơn, ít style hơn).

        Args:
            docx_input: Đường dẫn file (str) hoặc BytesIO.

        Returns:
            Chuỗi HTML đầy đủ.
        """
        temp_path = None
        start_t = time.time()
        try:
            if isinstance(docx_input, BytesIO):
                temp_path = _write_temp_file(docx_input.read(), ".docx")
                docx_file = temp_path
            elif isinstance(docx_input, str):
                if not os.path.exists(docx_input):
                    raise FileNotFoundError(f"File not found: {docx_input}")
                docx_file = docx_input
            else:
                raise ValueError("docx_input must be str or BytesIO")

            result       = mammoth.convert_to_html(docx_file)
            html_content = result.value
            output       = _MAMMOTH_HTML_TEMPLATE.format(html_content=html_content)

            logger.info("docx_to_html_mammoth_success", action="docx_to_html_mammoth", **{"event.duration": time.time()-start_t, "event.status": "success"}, chars=len(output))
            return output

        except Exception as e:
            logger.error("docx_to_html_mammoth_failed", action="docx_to_html_mammoth", **{"error.code": "IO", "error.message": str(e), "event.duration": time.time()-start_t, "event.status": "failure"}, exc_info=True)
            raise

        finally:
            if temp_path and os.path.exists(temp_path):
                os.remove(temp_path)


# ---------------------------------------------------------------------------
# Internal helpers
# ---------------------------------------------------------------------------

def _check_extension(file_name: str) -> None:
    ext = Path(file_name).suffix.lower()
    if ext not in ALLOWED_EXTENSIONS:
        raise ValueError(f"Invalid file format '{ext}'. Only {ALLOWED_EXTENSIONS} are allowed.")


def _check_size(file_name: str, size: int) -> None:
    if size > MAX_FILE_SIZE_BYTES:
        raise ValueError(f"File '{file_name}' exceeds 20MB limit ({size} bytes).")


def _write_temp_file(data: bytes, suffix: str) -> str:
    """Ghi bytes vào temp file và trả về đường dẫn."""
    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as f:
        f.write(data)
        return f.name


def _load_docx(source: Union[str, BufferedReader, BytesIO, FileStorage]) -> Document:
    """Load Document từ nhiều loại input khác nhau."""
    if isinstance(source, str):
        if not Path(source).exists():
            raise FileNotFoundError(f"File not found: {source}")
        return Document(source)
    elif isinstance(source, FileStorage):
        return Document(source.stream)
    elif isinstance(source, (BufferedReader, BytesIO)):
        return Document(source)
    else:
        raise ValueError(f"Unsupported input type: {type(source).__name__}")


def _get_alignment_style(para) -> str:
    mapping = {
        WD_ALIGN_PARAGRAPH.CENTER:  "text-align: center;",
        WD_ALIGN_PARAGRAPH.RIGHT:   "text-align: right;",
        WD_ALIGN_PARAGRAPH.JUSTIFY: "text-align: justify;",
    }
    return mapping.get(para.alignment, "text-align: left;")


def _render_runs(runs) -> str:
    """Render danh sách run thành HTML inline."""
    parts = []
    for run in runs:
        if not run.text:
            continue
        text = html.escape(run.text)

        # Inline formatting
        if run.bold:
            text = f"<strong>{text}</strong>"
        if run.italic:
            text = f"<em>{text}</em>"
        if run.underline:
            text = f"<u>{text}</u>"

        # Span styles
        styles = []
        if run.font.name:
            styles.append(f"font-family: '{run.font.name}';")
        if run.font.size:
            styles.append(f"font-size: {run.font.size.pt}pt;")
        try:
            if run.font.color.rgb:
                styles.append(f"color: #{run.font.color.rgb};")
        except Exception:
            pass

        if styles:
            text = f'<span style="{"".join(styles)}">{text}</span>'
        parts.append(text)

    return "".join(parts)


def _render_table(table) -> str:
    """Render một table thành HTML."""
    rows = ['<table border="1" style="border-collapse: collapse; width: 100%;">']
    for row in table.rows:
        cells = ["<tr>"]
        for cell in row.cells:
            cell_html = "".join(
                f"<p>{_render_runs(para.runs)}</p>"
                for para in cell.paragraphs
                if para.text.strip()
            )
            cells.append(f'<td style="padding: 5px;">{cell_html}</td>')
        cells.append("</tr>")
        rows.append("".join(cells))
    rows.append("</table>")
    return "\n".join(rows)


def _wrap_html(body: str) -> str:
    return f"""<!DOCTYPE html>
<html lang="vi">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Document</title>
    <style>
        body {{ font-family: Arial, sans-serif; max-width: 800px; margin: 20px auto; line-height: 1.6; color: #333; }}
        h1, h2, h3, h4, h5, h6 {{ color: #2c3e50; margin: 10px 0; }}
        p {{ margin: 8px 0; }}
        ul, ol {{ margin: 10px 0; padding-left: 20px; }}
        li {{ margin-bottom: 5px; }}
        table {{ border-collapse: collapse; width: 100%; margin: 10px 0; }}
        td, th {{ border: 1px solid #ddd; padding: 8px; }}
    </style>
</head>
<body>
{body}
</body>
</html>"""

if __name__ == "__main__":
    file_path = "/home/ubuntu/projects/AI/git/dev/v03/v03-document-management-services-dev/core/common/reader/12_NQ-CP_553515.docx"
    
    reader = DocumentProcessor()
    
    content = reader.read_docx(file_path)
    print(content)